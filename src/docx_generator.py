import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime

def generate_order_docx(order_data, filename):
    """
    Generate a DOCX for an order/invoice with the following layout:

    Header (centered):
      - Date line: e.g. "Sunday, 31 Mar 2025 @ 5PM" (16pt, bold)
      - Customer Name: (18pt, bold, italic)
      - Order Number: (12pt)
      - Total pax: [total_pax] + [difference between entrée and desserts] (16pt, bold)

    Then a larger gap before the sections.

    Sections:
      - Entrée (16pt, bold)
         Each item: "quantity ItemName (comment if applicable)" (14pt, bold, color #292929)
      - Mains (16pt, bold)
         ...
      - Desserts (16pt, bold)
         ...
    """
    doc = Document()

    # Set default style: Arial, 12pt.
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    def add_centered_line(text, font_size, bold, italic=False, space_after=6):
        p = doc.add_paragraph()
        p.alignment = 1  # center
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        return p

    # Header fields
    date_line = order_data.get("date", datetime.now().strftime("%A, %d %b %Y @ %I%p"))
    customer_name = order_data.get("customer_name", "[Customer Name]")
    order_number = order_data.get("order_number", "[order number]")
    total_pax = order_data.get("total_pax", 0)

    def sum_items(items):
        return sum(item[1] for item in items)
    total_entree = sum_items(order_data.get("entree", []))
    total_desserts = sum_items(order_data.get("desserts", []))
    diff = total_entree - total_desserts

    # Header paragraphs
    add_centered_line(date_line, font_size=16, bold=True, space_after=6)
    add_centered_line(customer_name, font_size=18, bold=True, italic=True, space_after=6)
    add_centered_line(order_number, font_size=12, bold=False, space_after=6)
    add_centered_line(f"Total pax: {total_pax} + {diff}", font_size=16, bold=True, space_after=20)

    # Section builder
    def add_section(title, items):
        add_centered_line(title, font_size=16, bold=True, space_after=6)
        for (item_name, qty, comment) in items:
            if qty > 0:
                line = f"{qty} {item_name}"
                if comment and comment.strip():
                    line += f" ({comment.strip()})"
                p = doc.add_paragraph()
                p.alignment = 1
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(line)
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x29, 0x29, 0x29)  # #292929
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_section("Entrée", order_data.get("entree", []))
    add_section("Mains", order_data.get("mains", []))
    add_section("Desserts", order_data.get("desserts", []))

    doc.save(filename)
    return filename

def save_order_as_docx(order_data, output_folder):
    """
    Create a DOCX file for the given order using the naming convention:
      ordernumber_customername_invoicenumber.docx
    """
    os.makedirs(output_folder, exist_ok=True)
    cust_name = order_data.get("customer_name", "Unknown")
    sanitized_name = re.sub(r'\W+', '', cust_name)
    order_num = order_data.get("order_number", "0000")
    invoice_num = order_data.get("invoice_number", order_num)
    filename = f"{order_num}_{sanitized_name}_{invoice_num}.docx"
    full_path = os.path.join(output_folder, filename)
    return generate_order_docx(order_data, full_path)

# For testing purposes:
if __name__ == "__main__":
    sample_order = {
        "order_number": "1234",
        "customer_name": "John Doe",
        "invoice_number": "INV-00001",
        "date": "Sunday, 31 Mar 2025 @ 5PM",
        "total_pax": 10,
        "entree": [("Salad", 5, "Fresh"), ("Soup", 0, "")],
        "mains": [("Steak", 3, ""), ("Fish", 2, "Grilled")],
        "desserts": [("Ice Cream", 4, "")],
    }
    output_folder = os.getcwd()  # or specify another folder
    docx_path = save_order_as_docx(sample_order, output_folder)
    print("DOCX generated at:", docx_path)
